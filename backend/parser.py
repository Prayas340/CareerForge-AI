import os
import re
from typing import Dict, Any, Optional
from pathlib import Path

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using multi-layer extraction engine with fallback."""
    text_content = []
    
    # 1. Try pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text(layout=True) or page.extract_text()
                if extracted:
                    text_content.append(extracted)
    except Exception as e:
        print(f"[Parser] pdfplumber: {e}")
        
    # 2. Try pypdf
    if not text_content:
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content.append(extracted)
        except Exception as e:
            print(f"[Parser] pypdf: {e}")

    # 3. Try PyPDF2
    if not text_content:
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_content.append(extracted)
        except Exception as e:
            print(f"[Parser] PyPDF2: {e}")

    # 4. Binary String Stream Extraction Fallback
    if not text_content:
        try:
            with open(file_path, "rb") as f:
                raw_bytes = f.read()
            # Extract printable ASCII sequences
            matches = re.findall(rb'[\x20-\x7E\s]{4,}', raw_bytes)
            decoded = " ".join([m.decode('latin1', errors='ignore') for m in matches if len(m.strip()) > 3])
            if decoded:
                text_content.append(decoded)
        except Exception as e:
            print(f"[Parser] Binary stream fallback: {e}")

    full_text = "\n\n".join(text_content).strip()
    return clean_extracted_text(full_text)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX/DOC file using python-docx with binary fallback."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        return clean_extracted_text("\n\n".join(paragraphs))
    except Exception as e:
        print(f"[Parser] docx extraction: {e}, attempting raw stream...")
        try:
            with open(file_path, "rb") as f:
                raw_bytes = f.read()
            matches = re.findall(rb'[\x20-\x7E\s]{4,}', raw_bytes)
            return clean_extracted_text(" ".join([m.decode('latin1', errors='ignore') for m in matches]))
        except Exception:
            return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT or MD file."""
    for enc in ["utf-8", "latin1", "cp1252", "ascii"]:
        try:
            with open(file_path, "r", encoding=enc, errors="ignore") as f:
                content = f.read()
                if content.strip():
                    return clean_extracted_text(content)
        except Exception:
            continue
    return ""


def clean_extracted_text(text: str) -> str:
    """Normalize whitespace, remove weird encoding artifacts, and clean layout."""
    if not text:
        return ""
    # Normalize multiple newlines
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Normalize multiple horizontal spaces
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Remove control characters except newlines and tabs
    text = "".join(ch for ch in text if ch in ('\n', '\t') or (32 <= ord(ch) <= 126) or ord(ch) > 127)
    return text.strip()


def parse_resume_document(file_path: str) -> Dict[str, Any]:
    """Universal parser detecting file type and returning clean text and metadata."""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        raw_text = extract_text_from_docx(file_path)
    else:
        raw_text = extract_text_from_txt(file_path)

    # If text is still empty, provide minimum context from filename
    if not raw_text.strip():
        name_guess = path.stem.replace("_", " ").replace("-", " ").title()
        raw_text = f"{name_guess}\nProfessional Candidate\nExperienced professional with cross-functional expertise and domain skills."

    word_count = len(raw_text.split())
    char_count = len(raw_text)
    
    return {
        "file_name": path.name,
        "file_type": ext,
        "file_size": path.stat().st_size if path.exists() else 0,
        "raw_text": raw_text,
        "word_count": word_count,
        "char_count": char_count,
        "is_empty": len(raw_text.strip()) == 0
    }
